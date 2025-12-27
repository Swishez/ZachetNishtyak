from django.shortcuts import render
from django.http import HttpResponse
from django.conf import settings
import os
from docx import Document
from django.views.decorators.csrf import csrf_exempt

@csrf_exempt  # или > используйте csrf_token, чтобы не было ошибок
def create_letter(request):
    if request.method == 'POST':
        address = request.POST.get('address_text', '')
        ownline = request.POST.get('ownline_text', '')

        # Создаем документ
        doc = Document()
        doc.add_paragraph(address)
        doc.add_paragraph('')
        doc.add_paragraph(ownline)

        # Создаем временный файл
        filename = 'письмо_{}.docx'.format(request.user.id if request.user.is_authenticated else 'anon')
        filepath = os.path.join(settings.MEDIA_ROOT, filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        doc.save(filepath)

        # Отправляем файл для скачивания
        with open(filepath, 'rb') as f:
            response = HttpResponse(f.read(),
                                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename={filename}'
        # Можно удалить файл после отправки, если хотите (не обязательно)
        return response
    else:
        # GET-запрос — возвращаем страницу с формой
        return render(request, 'create_letter.html')
