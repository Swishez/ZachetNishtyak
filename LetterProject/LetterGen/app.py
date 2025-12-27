from django.shortcuts import render
from django.http import FileResponse
import os
from django.conf import settings
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from lxml import etree
from docx.oxml.ns import qn
from docx.shared import Pt
import tempfile
from django.http import FileResponse
    
def index(request):
    return render(request, 'main/index.html')

def create_letter(request):
    # Создаем документ
    doc = Document()

    # Таблица с 2 ячейками
    table = doc.add_table(rows=1, cols=2)
    cell_left = table.cell(0, 0)
    cell_left._element.clear_content()
    cell_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    contact_info = [
        "Российская Федерация",
        "Автономная некоммерческая профессиональная образовательная организация",
        "«Кооперативный техникум Тамбовского облпотребсоюза»",
        "392020, г. Тамбов,",
        "ул. Пролетарская, д. 252/2",
        "Тел.:(4752) 53-53-70",
        "E-mail: tkt33@mail.ru",
        "от ____________г. №______",
        "На №_________от__________"
    ]

    for line in contact_info:
        p = cell_left.add_paragraph(line)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = p.runs[0]
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    cell_right = table.cell(0, 1)
    cell_right._element.clear_content()
    cell_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    address_text = "К директору\n"
    p_right = cell_right.add_paragraph(address_text)
    p_right.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_right = p_right.runs[0]
    run_right.font.size = Pt(10)

    # Убираем границы таблицы
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = etree.SubElement(tbl, qn('w:tblPr'))

    borders = etree.SubElement(tblPr, qn('w:tblBorders'))
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border_el = etree.SubElement(borders, qn(f'w:{border_name}'))
        border_el.set(qn('w:val'), 'nil')

    # Основной текст
    text_ownline = ["Уважаемая ФИО"]
    for line in text_ownline:
        p1 = doc.add_paragraph(line)
        run = p1.runs[0]
        run.font.size = Pt(14)
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    text_lines = [
        "АНПОО «Кооперативный техникум Тамбовского облпотребсоюза» (Далее - кооперативный техникум) - одно из старейших ведущих учебных заведений среднего профессионального образования, осуществляющее подготовку специалистов по очной форме обучения для системы потребительской кооперации, различных отраслей государственного хозяйства страны, сферы частного предпринимательства.",
        "В настоящее время кооперативный техникум осуществляет подготовку по следующим специальностям:",
        "- Банковское дело",
        "- Информационные системы и программирование",
        "- Коммерция",
        "- Поварское и кондитерское дело",
        "- Товароведение и экспертиза качества потребительских товаров",
        "- Финансы",
        "- Экономика и бухгалтерский учет",
        "- Юриспруденция",
        "Прошу вас оказать содействие в ознакомлении учеников с информацией о деятельности кооперативного техникума",
        "Исполняющий обязанности директора                                          Н. А. Земской",
    ]

    for line in text_lines:
        p2 = doc.add_paragraph(line)
        p2.style.font.size = Pt(12)

        # Генерируем файл и сохраняем на сервере
        filename = "оформленное_письмо.docx"
        filepath = os.path.join(settings.MEDIA_ROOT, filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        doc.save(filepath)

        # Отправляете файл пользователю для скачивания
        return FileResponse(open(filepath, 'rb'), as_attachment=True, filename=filename)
    else:
        # GET-запрос — возвращаем страницу с формой
        return render(request, 'LetterGen/create_letter.html')
  