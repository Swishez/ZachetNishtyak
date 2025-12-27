from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from lxml import etree
from docx.oxml.ns import qn
from docx.shared import Pt
from flask import Flask, request, render_template, send_from_directory
import os
import docx

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate():
    address_text = request.form['address_text']
    ownline_text = request.form['ownline_text']

    # Создаем документ
    doc = docx.Document()

    # Здесь ваш код для оформления документа
    # Например, добавляем параграфы
    doc.add_paragraph(address_text)
    doc.add_paragraph(ownline_text)

    # Убедитесь, что папка для сохранения есть
    folder_path = 'generated_files'
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

    # Имя файла
    filename = 'оформленное_письмо.docx'
    file_path = os.path.join(folder_path, filename)

    # Сохраняем документ
    doc.save(file_path)

    # Возвращаем ссылку для скачивания
    return render_template('download.html', filename=filename)

@app.route('/download/<filename>')
def download(filename):
    return send_from_directory('generated_files', filename)

if __name__ == '__main__':
    app.run(debug=True)

# Создаем новый документ
doc = Document()

# Создаем таблицу с 1 строкой и 2 столбцами
table = doc.add_table(rows=1, cols=2)

# Заполняем левую ячейку — контактная информация
cell_left = table.cell(0, 0)

# Очищаем содержание ячейки, если есть
cell_left._element.clear_content()

# Свойство вертикального выравнивания прописываем через direct API
cell_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Массив строк для левой части
contact_info = [
    "Российская Федерация",
    "Автономная некоммерческая профессиональная образовательная организация",
    "«Кооперативный техникум Тамбовского облпотребсоюза»",
    "392020, г. Тамбов,",
    "ул. Пролетарская, д. 252/2",
    "Тел.:(4752) 53-53-70",
    "E-mail: tkt33@mail.ru",
    "от ____________г. №______",
    "На №_________от__________"
]

# Добавляем каждую строку как отдельный параграф
for line in contact_info:
    p = cell_left.add_paragraph(line)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.runs[0]
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)

# Заполняем правую ячейку — адрес
cell_right = table.cell(0, 1)

# Очищаем содержание правой ячейки
cell_right._element.clear_content()

cell_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

address_text = "К директору\nМБОУ «название №номер»\nинициалы"
p_right = cell_right.add_paragraph(address_text)
p_right.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run_right = p_right.runs[0]
run_right.font.size = Pt(10)

# Убираем границы таблицы — делаем ее невидимой
tbl = table._tbl
tblPr = tbl.find(qn('w:tblPr'))
if tblPr is None:
    tblPr = etree.SubElement(tbl, qn('w:tblPr'))

borders = etree.SubElement(tblPr, qn('w:tblBorders'))
for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    border_el = etree.SubElement(borders, qn(f'w:{border_name}'))
    border_el.set(qn('w:val'), 'nil')

# Добавляем основной текст письма

text_ownline = [
    "Уважаемая ФИО"
]
for line in text_ownline:
    p1 = doc.add_paragraph(line)
    run = p1.runs[0]
    run.font.size = Pt(14)
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

text_lines = [
    "АНПОО «Кооперативный техникум Тамбовского облпотребсоюза» (Далее - кооперативный техникум) - одно из старейших ведущих учебных заведений среднего профессионального образования, осуществляющее подготовку специалистов по очной форме обучения для системы потребительской кооперации, различных отраслей государственного хозяйства страны, сферы частного предпринимательства.",
    "В настоящее время кооперативный техникум осуществляет подготовку по следующим специальностям:",
    "- Банковское дело",
    "- Информационные системы и программирование",
    "- Коммерция",
    "- Поварское и кондитерское дело",
    "- Товароведение и экспертиза качества потребительских товаров",
    "- Финансы",
    "- Экономика и бухгалтерский учет",
    "- Юриспруденция",
    "Прошу вас оказать содействие в ознакомлении учеников с информацией о деятельности кооперативного техникума",
    "Исполняющий обязанности директора                                          Н. А. Земской",
]

for line in text_lines:
    p2 = doc.add_paragraph(line)
    p2.style.font.size = Pt(12)

# Сохранение файла
doc.save("оформленное_письмо.docx")