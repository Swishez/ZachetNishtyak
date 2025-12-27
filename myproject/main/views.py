from django.shortcuts import render, redirect
from django.conf import settings
import os
import uuid
import docx

def index(request):
    return render(request, 'main/index.html')

def generate_and_save(request):
    if request.method == 'POST':
        address_text = request.POST.get('address_text')
        ownline_text = request.POST.get('ownline_text')
        
        # Создаём новый документ
        doc = docx.Document()
        doc.add_paragraph(address_text)
        doc.add_paragraph()
        doc.add_paragraph(ownline_text)
        
        # Генерируем уникальное имя файла
        filename = f"{uuid.uuid4()}.docx"
        filepath = os.path.join(settings.MEDIA_ROOT, filename)
        
        # Сохраняем файл
        doc.save(filepath)
        
        # Перенаправляем на страницу скачивания
        return redirect('download', filename=filename)
    else:
        return redirect('index')

def download(request, filename):
    from django.http import FileResponse
    filepath = os.path.join(settings.MEDIA_ROOT, filename)
    return FileResponse(open(filepath, 'rb'), as_attachment=True, filename=filename)
